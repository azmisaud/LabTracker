from datetime import datetime
from django.contrib.auth.forms import AuthenticationForm
from django.shortcuts import render, redirect
from django.contrib.auth import login as auth_login, logout
from django.views.decorators.csrf import csrf_exempt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from LabTrackerAMU import settings
from LabTrackerAMU.decorators import student_required
from faculty.models import LastDateOfWeek
from .forms import StudentSignUpForm, EnrollmentFacultyForm, DateOfBirthForm, PasswordResetForm
from problems.models import Problem, ProblemCompletion, WeekCommit
from django.http import HttpResponse, JsonResponse
from docx import Document
from io import BytesIO
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests
from PIL import Image
import json
from .models import Student, PasswordResetToken
from django.contrib.auth import authenticate, login as auth_login


def student_signup(request):
    """
    Handles the student sign-up view.

    - If the request method is POST, it processes the submitted form.
      If the form is valid, the user is saved and the page is redirected to the 'student_signup' view.
    - If the form is invalid, the form is re-rendered with validation errors.
    - If the request method is GET, an empty form is displayed for user input.

    Args:
        request (HttpRequest): The incoming HTTP request from the client.

    Returns:
        HttpResponse: Renders the student sign-up page with either an empty form or form with errors.
    """
    # Handle form submission if the request is a POST
    if request.method == 'POST':
        form = StudentSignUpForm(request.POST)

        # Check if the form data is valid
        if form.is_valid():
            # Save the user and create the student account
            user = form.save()
            # Redirect to the login page upon successful sign-up
            return redirect('student_login')
        else:
            # If the form is invalid, re-render the page with form errors
            return render(request, 'students/signup.html', {'form': form})

    # For GET requests, create an empty form
    else:
        form = StudentSignUpForm()

    # Render the sign-up template with the form
    return render(request, 'students/signup.html', {'form': form})

def student_login(request):
    """
    Handle the login process for students using their username (displayed as email in the UI).

    Args:
        request: HttpRequest object that contains metadata about the request.

    Returns:
        HttpResponseRedirect: Redirects to the 'student_dashboard' page upon successful login.
        HttpResponse: Renders the login page with an error if the credentials are invalid.
    """
    if request.method == 'POST':
        username = request.POST.get('email')  # Get the username from the "email" input field
        password = request.POST.get('password')

        # Authenticate the user
        user = authenticate(request, username=username, password=password)

        if user is not None:
            # If the user exists, log them in
            auth_login(request, user)
            return redirect('student_dashboard')
        else:
            # If authentication fails, render the form with an error
            error_message = "Invalid username or password."
            return render(request, 'students/login.html', {'error': error_message})

    return render(request, 'students/login.html')

def student_logout(request):
    """
    Logs out the current student and redirects them to the login page.

    This view function is responsible for logging out the currently authenticated student
    by calling Django's built-in `logout` function, which clears the session data.
    Once logged out, the user is redirected to the student login page.

    Args:
        request (HttpRequest): The incoming HTTP request object that triggers the logout.

    Returns:
        HttpResponseRedirect: A redirect to the 'student_login' view after the student is logged out.

    Example Usage:
        A student clicking the "logout" button will trigger this view. Upon execution:
        - The student will be logged out, clearing their session and authentication data.
        - The student will be redirected to the login page where they can log in again.

    Notes:
        - This function uses Django's `logout` function to handle the actual session clearing and user logout.
        - No special permissions are required to access this view since it is typically triggered by authenticated users.
    """

    # Call Django's built-in logout function to clear the user's session and authentication data
    logout(request)

    # Redirect the user to the student login page after logout
    return redirect('student_login')


@student_required
def student_dashboard(request):
    """
    Handles the student dashboard view. Provides an overview of the student's progress
    by retrieving problems assigned to their course and semester, their completion status,
    weekly commits, and deadlines. The overall progress is calculated as well as weekly
    progress.

    Parameters:
    request (HttpRequest): The HTTP request object which contains the logged-in student.

    Returns:
    HttpResponse: The rendered 'students/dashboard.html' template with the student's progress data.
    """

    # Retrieve the logged-in student and their course and semester information
    student = request.user
    course = student.course
    semester = student.semester

    # Fetch problems for the student's course and semester, ordered by week and problem number
    problems = Problem.objects.filter(course=course, semester=semester).order_by('week', 'problemNumber')

    # Fetch the student's completion records for the filtered problems
    problem_completions = ProblemCompletion.objects.filter(student=student, problem__in=problems).order_by('problem')

    # Fetch WeekCommit entries for the student, ordered by week number
    week_commits = WeekCommit.objects.filter(student=student).order_by('week_number')

    # Convert week commit data into a dictionary for easy access by week number
    week_commit_data = {week_commit.week_number: week_commit for week_commit in week_commits}

    # Fetch deadlines (last dates) for each week of the student's course and semester
    week_last_dates = LastDateOfWeek.objects.filter(course=course, semester=semester)

    # Store week deadlines in a dictionary for quick lookup by week number
    week_last_date_data = {week_last.week: week_last.last_date for week_last in week_last_dates}

    # Initialize dictionaries to track weekly and overall progress
    weekly_progress = {}
    overall_progress = {'total': len(problems), 'completed': 0}

    # List to store the completion status of each problem for rendering purposes
    completion_list = []

    # Loop through each problem and calculate the student's progress
    for prob in problems:
        week = prob.week

        # Initialize weekly progress if it hasn't been done for the current week
        if week not in weekly_progress:
            weekly_progress[week] = {
                'total': 0,
                'completed': 0,
                'last_commit_time': None,
                'last_commit_hash': None,
                'last_date': week_last_date_data.get(week, 'No Deadline')  # Provide deadline if available
            }

        # Increment the total number of problems for the current week
        weekly_progress[week]['total'] += 1

        # Check if the current problem is completed by the student
        completion = problem_completions.filter(problem=prob).first()
        is_completed = completion and completion.is_completed

        # Append the completion status for the current problem
        completion_list.append({
            'problem_id': prob.id,
            'is_completed': is_completed
        })

        # If the problem is completed, update weekly and overall progress
        if is_completed:
            weekly_progress[week]['completed'] += 1
            overall_progress['completed'] += 1

        # If commit data is available for the week, update the last commit information
        if week in week_commit_data:
            week_commit = week_commit_data[week]
            weekly_progress[week]['last_commit_time'] = week_commit.last_commit_time
            weekly_progress[week]['last_commit_hash'] = week_commit.last_commit_hash

    # Calculate the overall completion percentage (avoid division by zero)
    overall_progress['percentage'] = (
        (overall_progress['completed'] / overall_progress['total']) * 100
        if overall_progress['total'] > 0
        else 0
    )

    # Prepare the context for rendering the template
    context = {
        'student': student,
        'weekly_progress': weekly_progress,  # Progress for each week
        'overall_progress': overall_progress,  # Overall progress across all weeks
        'problems': problems,  # List of problems for the course
        'completion_list': completion_list,  # Completion status for each problem
    }

    # Render and return the student dashboard template with the context data
    return render(request, 'students/dashboard.html', context)


@student_required
def generate_problem_doc(request):
    """
    Generates a downloadable Word document containing a structured list of problems for a student based on their course and semester.

    The document includes:
    - A title with the student's enrollment and faculty numbers in the footer.
    - A table listing the problems, organized by week and problem number, along with their descriptions and optional images.
    - Proper formatting for margins, alignment, and fonts.

    Functionality:
    1. Fetches all problems associated with the student's course and semester.
    2. Creates a Word document using python-docx.
    3. Sets margins and customizes the table layout.
    4. Merges table cells where appropriate for weeks with multiple problems.
    5. Handles images by inserting them in the corresponding table cell.
    6. Generates and returns the document as a downloadable file.

    Args:
        request: The HTTP request object. Expects the user to be logged in as a student.

    Returns:
        HttpResponse: A downloadable Word document (.docx) containing the problem list.
    """
    student = request.user
    course = student.course
    semester = student.semester

    # Fetch problems associated with the student's course and semester
    problems = Problem.objects.filter(course=course, semester=semester).order_by('week', 'problemNumber')

    # Initialize a Word document
    doc = Document()

    # Set margins
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Add footer with student details
    footer = section.footer
    left_paragraph = footer.add_paragraph()
    left_paragraph.text = f"{student.enrollment_number}         {student.first_name}   {student.last_name}             {student.faculty_number}"
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add document title
    title = doc.add_paragraph()
    title_run = title.add_run('INDEX')
    title_run.bold = True
    title_run.font.size = Pt(24)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Create table with a predefined layout
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Set column widths
    total_width = Cm(18).pt * 20
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(total_width)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    column_widths = [Cm(0.5), Cm(0.5), Cm(14), Cm(0.5), Cm(2.5)]

    # Apply column widths to the header row
    for row in table.rows:
        for idx, width in enumerate(column_widths):
            row.cells[idx].width = width

    # Set header row content
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'W\nE\nE\nK'
    hdr_cells[1].text = 'P\nR\nO\nB\nL\nE\nM'
    hdr_cells[2].text = 'PROBLEMS WITH DESCRIPTION'
    hdr_cells[3].text = 'P\nA\nG\nE'
    hdr_cells[4].text = 'SIGNATURE OF TEACHER WITH DATE'

    # Style the header row
    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(14)

    # Organize problems by week
    problems_by_week = {}
    for problem in problems:
        if problem.week not in problems_by_week:
            problems_by_week[problem.week] = []
        problems_by_week[problem.week].append(problem)

    # Add rows for each problem
    current_row_index = 1
    for week, problems_in_week in problems_by_week.items():
        num_problems = len(problems_in_week)

        for idx, problem in enumerate(problems_in_week):
            row_cells = table.add_row().cells
            row_cells[1].text = problem.problemNumber.replace('Problem ', '') + '#'
            row_cells[1].paragraphs[0].runs[0].bold = True
            row_cells[1].paragraphs[0].runs[0].font.size = Pt(11)
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Add problem description
            description_cell = row_cells[2]
            description_cell.text = problem.description
            for paragraph in description_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Comic Sans MS'
                    run.font.size = Pt(11)
            description_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            description_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Add problem image (with production URL handling)
            if problem.image:
                image_url = request.build_absolute_uri(problem.image.url)
                image_path = f"{settings.MEDIA_ROOT}/{image_url.split('/')[-1]}"
                docx_image = row_cells[2].add_paragraph().add_run().add_picture(image_path, width=Cm(7))

            row_cells[3].text = ""
            row_cells[4].text = ""

        # Merge cells for weeks with multiple problems
        if num_problems > 1:
            week_cell = table.cell(current_row_index, 0)
            week_cell.merge(table.cell(current_row_index + num_problems - 1, 0))
            week_cell.text = str(week)
            week_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            week_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            week_cell.paragraphs[0].runs[0].bold = True
            week_cell.paragraphs[0].runs[0].font.size = Pt(14)

            signature_cell = table.cell(current_row_index, 4)
            signature_cell.merge(table.cell(current_row_index + num_problems - 1, 4))
            signature_cell.text = ""
            signature_cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
            signature_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        else:
            table.cell(current_row_index, 0).text = str(week)
            table.cell(current_row_index, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(current_row_index, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.cell(current_row_index, 0).paragraphs[0].runs[0].bold = True
            table.cell(current_row_index, 0).paragraphs[0].runs[0].font.size = Pt(14)
            table.cell(current_row_index, 4).text = ""
            table.cell(current_row_index, 4).alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(current_row_index, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        current_row_index += num_problems

    # Save document to file stream
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # Return the Word document as a downloadable file
    response = HttpResponse(file_stream,
                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="index.docx"'

    return response

def fetch_url_content(url):
    """
    Fetches the content of a given URL using an HTTP GET request.

    Args:
        url (str): The URL from which to fetch the content.

    Returns:
        str: The content of the URL if the request is successful,
             or an error message if there is an issue with the request.

    Raises:
        None: All exceptions are caught and handled within the function.
    """
    try:
        response = requests.get(url)
        response.raise_for_status()
        return response.text
    except requests.exceptions.RequestException:
        return 'Error fetching URL'


def fetch_image(url):
    """
    Fetches an image from a given URL using an HTTP GET request.

    Args:
        url (str): The URL of the image to be fetched.

    Returns:
        tuple: A tuple containing:
            - image_data (BytesIO): The image data as a BytesIO stream.
            - image (PIL.Image): The image object from Pillow for processing.
    """

    try:
        response = requests.get(url)
        response.raise_for_status()
        image_data = BytesIO(response.content)
        image = Image.open(image_data)
        return image_data, image
    except requests.exceptions.RequestException:
        return None

@student_required
def generate_file(request, week):
    """
    Generates a Microsoft Word document (.docx) containing problem details, solutions, and output images
    for a given student, course, semester, and week. This document is then returned as a downloadable file.

    Args:
        request (HttpRequest): The HTTP request object, containing information about the logged-in student.
        week (int): The week number for which the problems and solutions are to be retrieved.

    Workflow:
        1. Retrieves the logged-in student and their respective course, semester, and the requested week's problems.
        2. Fetches problem details along with the student's completion status, solution URL, and output image.
        3. Generates a .docx file containing:
            - Problem descriptions
            - Solutions (if available) or a placeholder text if not
            - Output images (if available) or a placeholder message if not
        4. Adds a footer indicating the generation time of the document.
        5. Returns the document as an HTTP response for downloading.

    Returns:
        HttpResponse: A response containing the .docx file as an attachment to be downloaded.

    Notes:
        - If a problem's solution or output image is unavailable, a placeholder message is inserted.
        - The generated document uses "Comic Sans MS" for the solution text, and bold formatting for problem titles.
        - The margins of the document are customized for a more compact layout.

    Example:
        A student can download the document for Week 5 by visiting the corresponding URL, where the week parameter is provided.
        The document will contain all problems for that week, including their descriptions, solutions, and output images.

    Raises:
        None: All exceptions are handled internally, and placeholders are used if data (solution or image) is missing.
    """
    student = request.user
    course = student.course
    semester = student.semester

    # Fetching the problems for the specified week
    problems = Problem.objects.filter(course=course, semester=semester, week=week).order_by('problemNumber')
    problem_completions = ProblemCompletion.objects.filter(student=student, problem__in=problems)

    buffer = BytesIO()
    document = Document()

    # Setting margins for the document
    section = document.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # Adding the week heading
    heading = document.add_heading(level=1)
    heading_run=heading.add_run(f"Week {week}")
    heading_run.font.name='Arial Black'
    heading_run.font.size=Pt(20)
    heading_run.font.bold=True
    heading_run.font.color.rgb=RGBColor(0x00,0x00,0x00)
    heading.alignment = 1  # Center the heading

    #Add spacing after the heading
    heading_format=heading.paragraph_format
    heading_format.space_after=Pt(20)

    # Custom style for the solution text
    styles = document.styles
    comic_sans_style = styles.add_style('ComicSansStyle', 2)  # '2' stands for Paragraph style
    font = comic_sans_style.font
    font.name = 'Comic Sans MS'
    font.size = Pt(11)
    font.italic = True

    # Iterating over problems and adding content to the document
    for prob in problems:
        completion = problem_completions.filter(problem=prob).first()
        solution_url = completion.solution_url if completion else None
        output_image_url = completion.output_image_url if completion else None

        # Adding the problem description
        p = document.add_paragraph()
        run = p.add_run(f"{prob.problemNumber}: {prob.description}")
        run.bold = True

        # Adding the solution
        if solution_url:
            solution_content = fetch_url_content(solution_url)
            document.add_heading("Solution: ", level=3)
            k = document.add_paragraph()
            runk = k.add_run(solution_content)
            runk.style = 'ComicSansStyle'
        else:
            document.add_heading("Solution: ", level=3)
            k = document.add_paragraph()
            runk = k.add_run("No Solution Available")
            runk.style = 'ComicSansStyle'

        # Adding the output image
        if output_image_url:
            image_data,image = fetch_image(output_image_url)
            if image_data and image:
                document_width=section.page_width-section.left_margin-section.right_margin
                document_width_in_pixels=int(document_width / Inches(1)) * image.info['dpi'][0] //96

                image_width,image_height=image.size

                document.add_heading("Output Image: ", level=3)

                if image_width>document_width_in_pixels:
                    aspect_ratio=image_height/image_width
                    new_width=document_width
                    new_height=new_width*aspect_ratio
                    document.add_picture(image_data,width=document_width)
                else:
                    document.add_picture(image_data,width=Inches(image_width / 96))

                last_paragraph=document.paragraphs[-1]
                last_paragraph.alignment=1
            else:
                document.add_heading("Output Image: ", level=3)
                document.add_paragraph("No Output Image Available")
        else:
            document.add_heading("Output Image: ", level=3)
            document.add_paragraph("No Output Image Available")

    # Footer with student's details

    footer=section.footer
    footer_paragraph=footer.paragraphs[0]
    footer_paragraph.text=f"{student.first_name} {student.last_name} | {student.enrollment_number} | {student.faculty_number} | {student.course} | Semester:{student.semester}"
    footer_paragraph.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adding the generation time footer
    document.add_paragraph()
    generation_time = datetime.now().strftime("%d-%m-%Y %H:%M:%S")
    footer_text = f"Generation Time: {generation_time}"
    document.add_paragraph(footer_text, style='Normal')

    # Saving the document in the buffer and returning it as an HTTP response
    document.save(buffer)
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="Week{week}.docx"'

    return response


@csrf_exempt
def forgot_password(request):
    """
    Handle the forgot password process.

    This view function handles a POST request where a student provides their
    enrollment number and faculty number. If the provided data is valid and matches
    a student in the database, it generates a password reset token for that student.
    The student must then enter their date of birth to complete the process.

    Parameters:
    - request (HttpRequest): The HTTP request object. Expects JSON data in the body
      with the enrollment number and faculty number.

    Returns:
    - JsonResponse: A JSON response containing either:
      - 'status' of 'success' and a 'message' prompting for further information along
        with the generated token, or
      - 'status' of 'error' with an appropriate error message.
    """
    if request.method == 'POST':
        # Load and parse JSON data from the request body
        data = json.loads(request.body)

        # Validate the provided data using the form
        form = EnrollmentFacultyForm(data)
        if form.is_valid():
            # Extract cleaned data (enrollment number and faculty number) from the form
            enrollment_number = form.cleaned_data['enrollment_number']
            faculty_number = form.cleaned_data['faculty_number']
            print(enrollment_number, faculty_number)  # Debugging: Print the enrollment and faculty numbers

            try:
                # Attempt to find the student matching the provided enrollment and faculty numbers
                student = Student.objects.get(enrollment_number=enrollment_number, faculty_number=faculty_number)

                # If student is found, create a password reset token
                token = PasswordResetToken.objects.create(student=student)

                # Return success response with token
                return JsonResponse({
                    'status': 'success',
                    'message': 'Please enter your date of birth',
                    'token': str(token.token)
                })
            except Student.DoesNotExist:
                # Return error response if no student is found
                return JsonResponse({
                    'status': 'error',
                    'message': 'Incorrect enrollment number or faculty number'
                })

        # Return error response if form data is invalid
        return JsonResponse({
            'status': 'error',
            'message': 'Invalid form data.'
        })

@csrf_exempt
def verify_date_of_birth(request):
    """
    Handle the verification of a student's date of birth for password reset.

    This view is part of the password reset process. It verifies the student's date of birth,
    which is sent along with a token that uniquely identifies the reset request. If the
    date of birth and the token are valid, the user is allowed to proceed with the password reset.

    Request Type:
    - POST: The function only processes POST requests. Other request types will result in an error response.

    Request Payload:
    - JSON object containing the following:
      - token (str): A unique token for the password reset request.
      - date_of_birth (str): The student's date of birth in `YYYY-MM-DD` format.

    Process:
    1. The function first attempts to parse the incoming JSON data. If the data is invalid, it returns an error.
    2. The `DateOfBirthForm` is used to validate the date of birth input.
    3. The token provided in the request is retrieved from the `PasswordResetToken` model.
    4. If the token is valid and the date of birth matches the student's record, a success message is returned.
       Otherwise, appropriate error messages are sent based on the issue (invalid date of birth, expired token, etc.).

    Returns:
    - JsonResponse: Returns a JSON response with either:
      - 'status': 'success', 'message': 'Valid date of birth, proceed to reset the password.'
      - 'status': 'error', 'message': 'Invalid date of birth', 'Invalid token', 'Time Limit Exceeded', etc.

    Errors Handled:
    - Invalid JSON structure in the request.
    - Invalid or expired token.
    - Invalid date of birth.

    Example JSON Response:
    {
        'status': 'success',
        'message': 'Valid date of birth, proceed to reset the password.'
    }
    """
    if request.method == 'POST':
        try:
            # Parse JSON data from the request body
            data = json.loads(request.body)
        except json.JSONDecodeError:
            return JsonResponse({'status': 'error', 'message': 'Invalid date of birth'})

        form = DateOfBirthForm(data)
        token = data.get('token')

        if form.is_valid():
            date_of_birth = form.cleaned_data['date_of_birth']
            try:
                # Attempt to retrieve the token from the PasswordResetToken model
                token = PasswordResetToken.objects.get(token=token)

                # Check if the token is still valid (not expired)
                if not token.is_valid():
                    return JsonResponse({'status': 'error', 'message': 'Time Limit Exceeded'})

                student = token.student

                # Verify the date of birth
                if student.date_of_birth == date_of_birth:
                    return JsonResponse({'status': 'success', 'message': 'Valid date of birth, proceed to reset the password.'})
                else:
                    return JsonResponse({'status': 'error', 'message': 'Invalid date of birth'})
            except PasswordResetToken.DoesNotExist:
                return JsonResponse({'status': 'error', 'message': 'Invalid token or Time Limit Exceeded'})

        # Handle invalid form data
        return JsonResponse({'status': 'error', 'message': 'Invalid form data.'})

    # Handle non-POST request methods
    return JsonResponse({'status': 'error', 'message': 'Invalid request method'})

@csrf_exempt
def reset_password(request):
    """
    Handle the password reset process for students.

    This view is part of the password reset process. It allows students to reset their password
    by providing a valid token and new password. The token must still be valid (i.e., not expired)
    and the new password must match the confirmation password.

    Request Type:
    - POST: This view accepts only POST requests to process the password reset.

    Request Payload:
    - JSON object containing the following:
      - token (str): A unique token generated during the password reset request.
      - new_password (str): The student's new password.
      - confirm_password (str): Confirmation of the new password to ensure they match.

    Process:
    1. The function first attempts to parse the incoming JSON data.
    2. The `PasswordResetForm` is used to validate the input fields (new password and confirmation).
    3. If valid, it checks if the provided token exists and is still valid.
    4. If the token is valid and not expired, the student's password is reset and saved, and the token is deleted.
    5. If any validation fails (invalid token, mismatched passwords, or invalid form data), an appropriate error message is returned.

    Returns:
    - JsonResponse: A JSON response with one of the following:
      - 'status': 'success', 'message': 'Password Reset Successful'
      - 'status': 'error', 'message': Specific error messages, such as:
        - 'Invalid Token'
        - 'Passwords do not match'
        - 'Invalid form data'
        - 'Time Limit Exceeded'

    Example JSON Response:
    {
        'status': 'success',
        'message': 'Password Reset Successful'
    }

    Errors Handled:
    - Invalid or expired token.
    - Password mismatch between `new_password` and `confirm_password`.
    - Invalid form data (e.g., missing fields).
    """
    if request.method == 'POST':
        # Parse the incoming JSON data
        data = json.loads(request.body)

        # Create an instance of PasswordResetForm to validate the data
        form = PasswordResetForm(data)
        token = data.get('token')

        if form.is_valid():
            # Extract and compare the passwords
            new_password = form.cleaned_data['new_password']
            confirm_password = form.cleaned_data['confirm_password']

            if new_password == confirm_password:
                try:
                    # Retrieve the password reset token
                    token2 = PasswordResetToken.objects.get(token=token)

                    # Check if the token is still valid
                    if not token2.is_valid():
                        return JsonResponse({'status': 'error', 'message': 'Time Limit Exceeded'})

                    # Update the student's password
                    student = token2.student
                    student.set_password(new_password)
                    student.save()

                    # Delete the used token to prevent reuse
                    token2.delete()

                    return JsonResponse({'status': 'success', 'message': 'Password Reset Successful'})

                except PasswordResetToken.DoesNotExist:
                    return JsonResponse({'status': 'error', 'message': 'Invalid Token'})
            else:
                return JsonResponse({'status': 'error', 'message': 'Passwords do not match'})
        # Handle invalid form data
        return JsonResponse({'status': 'error', 'message': 'Invalid form data.'})

    # Return an error for any request method other than POST
    return JsonResponse({'status': 'error', 'message': 'Invalid request method'})

def password_reset(request):
    """
    Render the password reset page for students.

    This view serves the HTML page where students can initiate the password reset process by
    providing their enrollment number and faculty number. It does not handle any POST data or
    perform any logic related to resetting passwords, but simply renders the form.

    Template:
    - `students/forgotPassword.html`: This HTML template contains the password reset form for students.

    Request Type:
    - GET: This view accepts only GET requests, rendering the password reset page.

    Returns:
    - HttpResponse: Renders the 'forgotPassword.html' template.

    Example Usage:
    - A student navigates to the password reset page by clicking a "Forgot Password" link, and
      this view renders the form where they can submit their enrollment and faculty numbers.
    """
    return render(request, 'students/forgotPassword.html')